#! python3

import docx

doc = docx.Document('C:\\Code\\Python\\AutomateTheBoringStuffWithPython_NoStarchPress\\wordReading.docx')

print ('len(doc.paragraphs): ', len(doc.paragraphs))
print('doc.paragraphs[0].text: ', doc.paragraphs[0].text)
print('doc.paragraphs[1].text: ', doc.paragraphs[1].text)
print('len(doc.paragraphs[1].runs): ', len(doc.paragraphs[1].runs))
print('doc.paragraphs[1].runs[0].text: ', doc.paragraphs[1].runs[0].text)
print('doc.paragraphs[1].runs[1].text: ', doc.paragraphs[1].runs[1].text)
print('doc.paragraphs[1].runs[2].text: ', doc.paragraphs[1].runs[2].text)
print('doc.paragraphs[1].runs[3].text: ', doc.paragraphs[1].runs[3].text)
print('doc.paragraphs[1].runs[4].text: ', doc.paragraphs[1].runs[4].text)