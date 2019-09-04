import docx

doc = docx.Document()
doc.add_paragraph('Hello World!')

paraObj1 = doc.add_paragraph('This is a second paragraph.')
paraObj2 = doc.add_paragraph('This is yet another paragraph.')

paraObj1.add_run(' This text is being added to the second paragraph.')

doc.save('C:\\Code\\Python\\AutomateTheBoringStuffWithPython_NoStarchPress\\multipleParagraphs.docx')