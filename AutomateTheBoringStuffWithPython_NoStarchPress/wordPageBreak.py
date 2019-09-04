import docx

doc = docx.Document()
doc.add_paragraph('This is the first page!')
doc.paragraphs[0].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
doc.add_paragraph('This is the second page!')
doc.save('C:\\Code\\Python\\AutomateTheBoringStuffWithPython_NoStarchPress\\twoPage.docx')